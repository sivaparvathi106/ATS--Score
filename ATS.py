import streamlit as st
import pdfplumber
from sentence_transformers import SentenceTransformer, util
import re
from sklearn.feature_extraction.text import ENGLISH_STOP_WORDS
from docx import Document
import boto3
import json



model = SentenceTransformer('all-MiniLM-L6-v2')

 

def extract_text_from_pdf(upload_file):
   text = ""
   if upload_file is not None:
     with pdfplumber.open(upload_file) as pdf:
        for page in pdf.pages:
           page_text = page.extract_text()
           if page_text:
              text += page_text + ""
   
   return text.strip()
   
def extract_text_from_doc(upload_file):
   text = ""
   doc = docx.Document(upload_file)
   for para in doc.paragraphs:
      text += para.text + " "
   return text.strip()
   
   
def save_document(text , filename="Rewritten_Resume.docx"):
   doc = Document()
   for line in text.split("\n"):
       if line.strip():
         doc.add_paragraph(line.strip())
   doc.save(filename)
   return filename
   
   
def extract_text_from_file(upload_file):
   if upload_file.name.endswith(".pdf"):
      return extract_text_from_pdf(upload_file)
   elif upload_file.name.endswith(".docx"):
      return extract_text_from_doc(upload_file)
   else:
      return ""

def extract_words_dynamic(text: str):
    text = text.lower()
    words = re.findall(r'\b[a-zA-Z]{3,}\b' , text)
    keywords =[ w for w in words if w not in ENGLISH_STOP_WORDS]
    
    return list(set(keywords))
    

          

st.set_page_config(page_title='AI ATS Resume scorer' , layout='centered')
st.title('AI ATS Resume Scorer')
st.write('Upload a Job Description and Resume To Check ATS Score')

jd_file = st.file_uploader('Upload Job Description (PDF/DOCX)' , type=['pdf' , 'docx'])
resume_files = st.file_uploader('Upload Resume (PDF/DOCX)' , type=['pdf' , 'docx'] , accept_multiple_files = True)

if st.button("Get ATS Score"):
  
  if jd_file and resume_files:
  
    jd_text = extract_text_from_file(jd_file)
    if not jd_text:
      st.error('could not extract from JD')
    else:
       jd_embeddings = model.encode(jd_text , convert_to_tensor=True)
       jd_words = extract_words_dynamic(jd_text)
       results =[]
       for resume in resume_files:
          resume_text = extract_text_from_file(resume)
          if not resume_text:
             continue
          resume_embeddings = model.encode(resume_text , convert_to_tensor=True)
          score = util.pytorch_cos_sim(jd_embeddings ,resume_embeddings).item() *100
          
          resume_keywords =resume_text.lower()
          
          matched = [kw for kw in jd_words if kw in resume_keywords ]
          
          missing = [kw for kw in jd_words if kw not in resume_keywords]
    
          results.append({
              "resume_name":resume.name,
              "score":score,
              "matched":matched,
              "missing":missing,
              "resume_text": resume_text
           })
       
       results = sorted(results , key = lambda x : x["score"] ,reverse=True)
       
       for i,r in enumerate(results , 1):
        
         st.write(f"**{i}. {r['resume_name']}** - ATS Score: {r['score']:.2f}%")
         st.write("✅ Matched Keywords:", ", ".join(r['matched']) if r['matched'] else "None")
         st.write("❌ Missing Keywords:", ", ".join(r['missing']) if r['missing'] else "None")
         
         
         st.write("---")
       
else:
   st.warning('Please upload both JD and Resume.')
   


    
    
  
    